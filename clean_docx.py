from docx import Document
from bs4 import BeautifulSoup
import os
from textblob import TextBlob
import re

input_dir = 
output_dir = 

def clean_and_save_docx(input_dir, output_dir):
    for filename in os.listdir(input_dir):
        if filename.endswith('.docx'):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, 'cleaned_' + filename)
            
            # Load the .docx file
            document = Document(input_path)

            # Create a new Document
            new_document = Document()

            # Iterate over the paragraphs in the document
            for paragraph in document.paragraphs: 
                text = paragraph.text

                # Clean the text as before...
                text = re.sub(r'[^\w\s]', '', text)
                text = re.sub(r'http\S+|www\S+', '', text)
                text = str(TextBlob(text).correct())
                soup = BeautifulSoup(text, 'html.parser')
                text = soup.get_text()
                text = '  '.join(text.split())

                # Add the cleaned text to the new document
                new_document.add_paragraph(text)

            # Iterate over the tables in the document
            for table in document.tables:
                # Create a new table in the new document
                new_table = new_document.add_table(rows=0, cols=len(table.columns))

                # Iterate over the rows in the table
                for row in table.rows:
                    cells = row.cells

                    # Add a new row to the new table
                    new_cells = new_table.add_row().cells

                    # Iterate over the cells in the row
                    for i, cell in enumerate(cells):
                        # Copy the text from the old cell to the new cell
                        new_cells[i].text = cell.text

            # Save the new document
            new_document.save(output_path)

clean_and_save_docx(input_dir, output_dir)